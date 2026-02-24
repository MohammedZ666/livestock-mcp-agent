import os
import argparse
from docx import Document
import textract
import unicodeconverter as uc
from pdfminer.high_level import extract_text

import re

def clean_xml_text(text):
    """
    Remove characters not allowed in DOCX XML
    """
    return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F]', '', text)

def looks_like_bijoy(text):
    """
    Detect common Bijoy patterns
    """
    patterns = ["¤", "‡", "†", "Ð", "cÖ", "m¤ú", "`ß"]
    return any(p in text for p in patterns)


def bijoy_to_unicode(text):
    if not text:
        return text

    if looks_like_bijoy(text):
        try:
            return uc.convert_bijoy_to_unicode(text)
        except Exception:
            return text

    return text


def convert_docx(input_path, output_path):
    """
    Convert DOCX containing Bijoy text to Unicode
    while preserving tables
    """

    doc = Document(input_path)

    # convert paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            para.text = bijoy_to_unicode(para.text)

    # convert tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        para.text = bijoy_to_unicode(para.text)

    doc.save(output_path)


def convert_doc(input_path, output_path):
    """
    Extract text from old DOC and convert
    """

    raw = textract.process(input_path)

    # preserve byte mapping
    text = raw.decode("latin-1", errors="ignore")

    unicode_text = bijoy_to_unicode(text)

    doc = Document()
    for line in unicode_text.split("\n"):
        doc.add_paragraph(line)

    doc.save(output_path)


def convert_pdf(input_path, output_path):

    text = extract_text(input_path)

    unicode_text = bijoy_to_unicode(text)

    # remove invalid XML characters
    unicode_text = clean_xml_text(unicode_text)

    doc = Document()

    for line in unicode_text.split("\n"):
        doc.add_paragraph(line)

    doc.save(output_path)


def process_file(filepath, output_dir):

    filename = os.path.basename(filepath)
    name, ext = os.path.splitext(filename)

    output_path = os.path.join(output_dir, name + "_unicode.docx")

    ext = ext.lower()

    if ext == ".docx":
        convert_docx(filepath, output_path)

    elif ext == ".doc":
        convert_doc(filepath, output_path)

    elif ext == ".pdf":
        convert_pdf(filepath, output_path)

    else:
        print(f"Skipping unsupported file: {filename}")
        return

    print(f"Converted: {filename} -> {output_path}")


def main():

    parser = argparse.ArgumentParser(
        description="Convert Bijoy PDF/DOC/DOCX to Unicode DOCX"
    )

    parser.add_argument(
        "files",
        nargs="+",
        help="Input files (.pdf .doc .docx)"
    )

    parser.add_argument(
        "-o",
        "--output",
        required=True,
        help="Output directory"
    )

    args = parser.parse_args()

    os.makedirs(args.output, exist_ok=True)

    for filepath in args.files:
        try:
            process_file(filepath, args.output)
        except Exception as e:
            print(f"Failed: {filepath} -> {e}")


if __name__ == "__main__":
    main()